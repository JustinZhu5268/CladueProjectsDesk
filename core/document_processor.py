"""Document text extraction for various file formats."""
from __future__ import annotations

import os
import uuid
import shutil
import logging
from pathlib import Path

from data.database import db
from config import DOCS_DIR

log = logging.getLogger(__name__)

# Supported extensions grouped by extraction method
TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".tsv", ".json", ".xml", ".yaml", ".yml",
                   ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h", ".hpp",
                   ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".scala",
                   ".sh", ".bash", ".sql", ".r", ".m", ".lua", ".toml", ".ini",
                   ".cfg", ".conf", ".log", ".html", ".css", ".jsx", ".tsx", ".vue"}


class DocumentProcessor:
    """Extracts text content from documents and manages project files."""

    def add_document(self, project_id: str, source_path: str) -> dict:
        """Add a document to a project. Returns document record dict."""
        src = Path(source_path)
        if not src.exists():
            raise FileNotFoundError(f"File not found: {source_path}")

        doc_id = str(uuid.uuid4())
        ext = src.suffix.lower()
        filename = src.name
        file_size = src.stat().st_size

        # Copy to project document storage
        dest_dir = DOCS_DIR / project_id
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_path = dest_dir / f"{doc_id}{ext}"
        shutil.copy2(str(src), str(dest_path))
        log.info("Copied document '%s' (%d bytes) to %s", filename, file_size, dest_path)

        # Extract text
        extracted = self._extract_text(dest_path, ext)
        token_count = self._estimate_tokens(extracted)

        # Save to database
        rel_path = str(dest_path.relative_to(DOCS_DIR))
        db.execute(
            """INSERT INTO documents
               (id, project_id, filename, file_path, extracted_text,
                token_count, file_type, file_size, created_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))""",
            (doc_id, project_id, filename, rel_path, extracted,
             token_count, ext, file_size),
        )
        log.info("Indexed document '%s': %d tokens extracted", filename, token_count)
        return {
            "id": doc_id, "filename": filename, "token_count": token_count,
            "file_type": ext, "file_size": file_size,
        }

    def remove_document(self, doc_id: str) -> None:
        """Remove a document from the project."""
        row = db.execute_one("SELECT file_path FROM documents WHERE id = ?", (doc_id,))
        if row:
            fp = DOCS_DIR / row["file_path"]
            if fp.exists():
                fp.unlink()
                log.debug("Deleted file %s", fp)
        db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
        log.info("Removed document %s", doc_id[:8])

    def get_project_documents(self, project_id: str) -> list[dict]:
        """Get all documents for a project."""
        rows = db.execute(
            "SELECT * FROM documents WHERE project_id = ? ORDER BY created_at ASC",
            (project_id,),
        )
        return [dict(r) for r in rows]

    def get_project_context(self, project_id: str) -> str:
        """Get concatenated document text for API context."""
        rows = db.execute(
            "SELECT filename, extracted_text FROM documents WHERE project_id = ? ORDER BY created_at ASC",
            (project_id,),
        )
        if not rows:
            return ""
        parts = []
        for r in rows:
            parts.append(f'<document name="{r["filename"]}">\n{r["extracted_text"]}\n</document>')
        return "\n\n".join(parts)

    def get_total_tokens(self, project_id: str) -> int:
        """Get total token count for all project documents."""
        row = db.execute_one(
            "SELECT COALESCE(SUM(token_count), 0) as total FROM documents WHERE project_id = ?",
            (project_id,),
        )
        return row["total"] if row else 0

    # ── Extraction Methods ─────────────────────────────

    def _extract_text(self, path: Path, ext: str) -> str:
        """Route extraction based on file type."""
        try:
            if ext == ".pdf":
                return self._extract_pdf(path)
            elif ext == ".docx":
                return self._extract_docx(path)
            elif ext in (".xlsx", ".xls"):
                return self._extract_xlsx(path)
            elif ext in TEXT_EXTENSIONS:
                return self._extract_text_file(path)
            else:
                return self._extract_text_file(path)  # Try as text
        except Exception:
            log.exception("Failed to extract text from %s", path.name)
            return f"[Error: Could not extract text from {path.name}]"

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    pages.append(f"--- Page {i+1} ---\n{text.strip()}")
            doc.close()
            result = "\n\n".join(pages)
            if not result.strip():
                log.warning("PDF '%s' appears to be scanned (no text extracted)", path.name)
                return "[This PDF appears to be a scanned image. Text extraction unavailable.]"
            return result
        except ImportError:
            log.warning("PyMuPDF not installed, trying pdfplumber")
            return self._extract_pdf_fallback(path)

    def _extract_pdf_fallback(self, path: Path) -> str:
        """Fallback PDF extraction using pdfplumber."""
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(f"--- Page {i+1} ---\n{text.strip()}")
            return "\n\n".join(pages) or "[No text could be extracted from this PDF]"
        except ImportError:
            return "[Error: No PDF library available. Install PyMuPDF or pdfplumber.]"

    def _extract_docx(self, path: Path) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n\n".join(paragraphs)
        except ImportError:
            return "[Error: python-docx not installed]"

    def _extract_xlsx(self, path: Path) -> str:
        """Extract text from Excel files."""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path), data_only=True, read_only=True)
            sheets = []
            for ws in wb.worksheets:
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows.append(",".join(cells))
                if rows:
                    sheets.append(f"--- Sheet: {ws.title} ---\n" + "\n".join(rows))
            wb.close()
            return "\n\n".join(sheets)
        except ImportError:
            return "[Error: openpyxl not installed]"

    def _extract_text_file(self, path: Path) -> str:
        """Extract text from plain text files with encoding detection."""
        try:
            # Try UTF-8 first
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            try:
                import chardet
                raw = path.read_bytes()
                detected = chardet.detect(raw)
                encoding = detected.get("encoding", "utf-8") or "utf-8"
                log.debug("Detected encoding %s for %s", encoding, path.name)
                return raw.decode(encoding, errors="replace")
            except ImportError:
                return path.read_text(encoding="utf-8", errors="replace")

    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count. ~4 chars per token as rough estimate."""
        try:
            import tiktoken
            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except ImportError:
            return len(text) // 4
